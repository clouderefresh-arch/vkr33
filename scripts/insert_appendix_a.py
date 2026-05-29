"""
Скрипт вставки страниц PDF "бух отчетность КИПИШУЗ.pdf" в качестве изображений
в раздел "ПРИЛОЖЕНИЕ А" документа ВКР.

- Каждая страница PDF становится отдельным изображением (Рисунок А.N)
- Подписи академического стиля на основании содержимого ВКР
- Старое содержимое Приложения А (агрегированный баланс) удаляется,
  а вместо него вставляется официальная отчётность ООО «КУПИШУЗ» за 2023–2025 годы
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


WORKSPACE = Path("/workspace")
SRC_DOCX = WORKSPACE / "source" / "работа" / "ВКР_Купишуз_работай по ней.docx"
OUT_DOCX = WORKSPACE / "output" / "ВКР_Купишуз_работай по ней.docx"
IMG_DIR = WORKSPACE / "images"

CAPTIONS = [
    "Рисунок А.1 — Сведения о юридическом лице ООО «КУПИШУЗ» из "
    "Государственного информационного ресурса бухгалтерской (финансовой) "
    "отчётности (титульная страница выгрузки)",

    "Рисунок А.2 — Бухгалтерский баланс ООО «КУПИШУЗ» на 31 декабря 2025 года "
    "(с сопоставимыми данными за 2023 и 2024 годы), лист 1: актив, "
    "разделы I и II",

    "Рисунок А.3 — Бухгалтерский баланс ООО «КУПИШУЗ» на 31 декабря 2025 года "
    "(с сопоставимыми данными за 2023 и 2024 годы), лист 2: актив "
    "(окончание) и пассив, разделы III–V",

    "Рисунок А.4 — Бухгалтерский баланс ООО «КУПИШУЗ» на 31 декабря 2025 года "
    "(с сопоставимыми данными за 2023 и 2024 годы), лист 3: пассив "
    "(окончание), пояснения и подпись уполномоченного представителя",

    "Рисунок А.5 — Отчёт о финансовых результатах ООО «КУПИШУЗ» за 2025 год "
    "(с сопоставимыми данными за 2024 год), лист 1: выручка, себестоимость "
    "продаж и прибыль от продаж",

    "Рисунок А.6 — Отчёт о финансовых результатах ООО «КУПИШУЗ» за 2025 год "
    "(с сопоставимыми данными за 2024 год), лист 2: чистая прибыль "
    "и подпись уполномоченного представителя",

    "Рисунок А.7 — Отчёт об изменениях капитала ООО «КУПИШУЗ» за 2025 год, "
    "лист 1: движение капитала за 2024 год",

    "Рисунок А.8 — Отчёт об изменениях капитала ООО «КУПИШУЗ» за 2025 год, "
    "лист 2: движение капитала за 2025 год",

    "Рисунок А.9 — Отчёт об изменениях капитала ООО «КУПИШУЗ» за 2025 год, "
    "лист 3: корректировки в связи с изменением учётной политики и "
    "чистые активы",

    "Рисунок А.10 — Отчёт о движении денежных средств ООО «КУПИШУЗ» "
    "за 2025 год (с сопоставимыми данными за 2024 год), лист 1: денежные "
    "потоки от текущих операций (поступления)",

    "Рисунок А.11 — Отчёт о движении денежных средств ООО «КУПИШУЗ» "
    "за 2025 год (с сопоставимыми данными за 2024 год), лист 2: денежные "
    "потоки от инвестиционных и финансовых операций",

    "Рисунок А.12 — Отчёт о движении денежных средств ООО «КУПИШУЗ» "
    "за 2025 год (с сопоставимыми данными за 2024 год), лист 3: сальдо "
    "денежных потоков и подпись уполномоченного представителя",
]

INTRO_TEXT = (
    "Бухгалтерская (финансовая) отчётность ООО «КУПИШУЗ» за 2025 год "
    "(с сопоставимыми данными за 2023 и 2024 годы) приводится в настоящем "
    "приложении в виде сканов страниц официальной выгрузки из Государственного "
    "информационного ресурса бухгалтерской (финансовой) отчётности (Ресурса БФО) "
    "Федеральной налоговой службы Российской Федерации. Отчётность включает "
    "бухгалтерский баланс, отчёт о финансовых результатах, отчёт об изменениях "
    "капитала и отчёт о движении денежных средств. Указанные формы являются "
    "источником данных для финансово-экономического анализа предприятия "
    "и обоснования параметров инвестиционного проекта, представленных "
    "в основной части выпускной квалификационной работы."
)

NOTE_TEXT = (
    "Источник: выгрузка Государственного информационного ресурса "
    "бухгалтерской (финансовой) отчётности (Ресурса БФО) Федеральной "
    "налоговой службы Российской Федерации, ИНН 7705935687, "
    "по состоянию на 28.05.2026."
)


def get_text(elem):
    return "".join(t.text or "" for t in elem.iter(qn("w:t")))


def find_heading_index(body, text):
    for i, ch in enumerate(body):
        if ch.tag == qn("w:p") and get_text(ch).strip() == text:
            return i
    raise ValueError(f"Heading not found: {text}")


def make_centered_paragraph(doc, text, *, font_size=None, italic=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            from docx.oxml import OxmlElement
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
        if font_size is not None:
            run.font.size = Pt(font_size)
        if italic:
            run.font.italic = True
        if bold:
            run.font.bold = True
    return p


def make_justified_paragraph(doc, text, *, font_size=None, indent_first_line_cm=1.25,
                             italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first_line_cm:
        p.paragraph_format.first_line_indent = Cm(indent_first_line_cm)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        rPr = run._element.get_or_add_rPr()
        from docx.oxml import OxmlElement
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
        if font_size is not None:
            run.font.size = Pt(font_size)
        if italic:
            run.font.italic = True
    return p


def insert_picture_paragraph(doc, image_path, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


def main():
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SRC_DOCX))
    body = doc.element.body

    start_idx = find_heading_index(body, "ПРИЛОЖЕНИЕ А")
    end_idx = find_heading_index(body, "ПРИЛОЖЕНИЕ Б")

    print(f"ПРИЛОЖЕНИЕ А @ body[{start_idx}], ПРИЛОЖЕНИЕ Б @ body[{end_idx}]")

    # Удаляем существующее содержимое Приложения А (между двумя заголовками,
    # не трогая сами заголовки).
    for i in range(end_idx - 1, start_idx, -1):
        body.remove(body[i])
    print(f"Удалено элементов между заголовками: {end_idx - start_idx - 1}")

    # Создаём новые элементы в конце документа (python-docx добавляет их
    # перед sectPr), а затем переносим их в нужное место сразу после
    # заголовка "ПРИЛОЖЕНИЕ А".
    new_elements = []

    # Пустой параграф-разделитель
    sep = doc.add_paragraph()
    new_elements.append(sep._element)

    # Вступительный текст
    intro = make_justified_paragraph(doc, INTRO_TEXT, font_size=14)
    new_elements.append(intro._element)

    # Изображения и подписи
    image_files = sorted(IMG_DIR.glob("page-*.png"))
    if len(image_files) != len(CAPTIONS):
        raise RuntimeError(
            f"Несовпадение: страниц {len(image_files)}, подписей {len(CAPTIONS)}"
        )
    for img, caption in zip(image_files, CAPTIONS):
        pic_p = insert_picture_paragraph(doc, img, width_cm=15.5)
        new_elements.append(pic_p._element)
        cap_p = make_centered_paragraph(doc, caption, font_size=14)
        new_elements.append(cap_p._element)
        # Доп. пустой параграф для воздуха
        empty = doc.add_paragraph()
        new_elements.append(empty._element)

    # Финальное примечание
    note = make_justified_paragraph(
        doc, NOTE_TEXT, font_size=12, italic=True, indent_first_line_cm=0
    )
    new_elements.append(note._element)

    # Извлекаем эти элементы из конца тела и вставляем после заголовка А
    for e in new_elements:
        body.remove(e)

    # Перепроверяем индекс заголовка А (мог не сместиться, но на всякий случай)
    start_idx = find_heading_index(body, "ПРИЛОЖЕНИЕ А")
    insert_pos = start_idx + 1
    for e in new_elements:
        body.insert(insert_pos, e)
        insert_pos += 1

    doc.save(str(OUT_DOCX))
    print(f"Сохранено: {OUT_DOCX}")


if __name__ == "__main__":
    main()
